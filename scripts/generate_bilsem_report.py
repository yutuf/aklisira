#!/usr/bin/env python3
"""Generate expanded BILSEM project report as DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

from report_content import (
    ABSTRACT_TEXT,
    BOLUM1,
    BOLUM2_LITERATUR,
    BOLUM3_YONTEM,
    BOLUM4_BULGULAR,
    BOLUM5_SONUC,
    BULGU_TABLE_CSV,
    BULGU_TABLE_DEMO,
    IS_ZAMAN_ROWS,
    KAYNAKLAR,
    OZ_TEXT,
    TEKNO_TABLE,
)


def set_margins(section):
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.75)
    section.right_margin = Cm(2.75)


def style_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)


def add_center_bold(doc, text, size=14, space_before=0, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_body(doc, text, single=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE if single else WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0 if single else 1.5
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_section_blocks(doc, sections):
    for heading, paragraphs in sections:
        add_heading2(doc, heading)
        for para in paragraphs:
            add_body(doc, para)


def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.bold = True
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_toc_entry(doc, title, page_hint, indent_cm=0):
    """Manual TOC line with dot leader and page hint (BİLSEM: ön bölüm Romen, gövde Arap)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_after = Pt(2)
    if indent_cm:
        pf.left_indent = Cm(indent_cm)
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(f"{title}\t{page_hint}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_table_of_contents(doc):
    # BİLSEM: ön bölüm sayfaları Romen rakamları (i, ii, …); BÖLÜM 1'den itibaren Arap rakamları.
    # Sayfa ipuçları Word'de içerik uzunluğuna göre güncellenmelidir; alan kodu yerine manuel TOC.
    add_center_bold(doc, "İÇİNDEKİLER", 14, 0, 18)
    add_body(doc, "Ön bölüm sayfa numaraları Romen rakamları ile; ana metin Arap rakamları ile gösterilmiştir.", single=True)
    doc.add_paragraph()
    add_toc_entry(doc, "İÇ KAPAK", "i")
    add_toc_entry(doc, "ETİK İLKELERE UYGUNLUK BEYANI", "ii")
    add_toc_entry(doc, "JÜRİ ONAY SAYFASI", "iii")
    add_toc_entry(doc, "ÖZ", "iv")
    add_toc_entry(doc, "ABSTRACT", "v")
    add_toc_entry(doc, "İÇİNDEKİLER", "vi")
    add_toc_entry(doc, "SİMGELER VE KISALTMALAR LİSTESİ", "vii")
    doc.add_paragraph()
    add_toc_entry(doc, "BÖLÜM 1: GİRİŞ", "1")
    add_toc_entry(doc, "1.1. Problem Durumu", "1", indent_cm=0.75)
    add_toc_entry(doc, "1.2. Amaç", "2", indent_cm=0.75)
    add_toc_entry(doc, "1.3. Önemi", "2", indent_cm=0.75)
    add_toc_entry(doc, "1.4. Araştırma Soruları", "3", indent_cm=0.75)
    add_toc_entry(doc, "1.5. Sınırlılıklar", "3", indent_cm=0.75)
    add_toc_entry(doc, "1.6. Varsayımlar", "4", indent_cm=0.75)
    add_toc_entry(doc, "BÖLÜM 2: İLGİLİ ARAŞTIRMALAR", "5")
    add_toc_entry(doc, "2.1. Oturma Düzeninin Öğrenmeye Etkisi", "5", indent_cm=0.75)
    add_toc_entry(doc, "2.2. Akran Öğrenmesi ve Akademik Eşleştirme", "6", indent_cm=0.75)
    add_toc_entry(doc, "2.3. Davranış Yönetimi ve Stratejik Yerleşim", "6", indent_cm=0.75)
    add_toc_entry(doc, "2.4. Optimizasyon Teknikleri ve Eğitimde Uygulanması", "7", indent_cm=0.75)
    add_toc_entry(doc, "2.5. Literatürdeki Boşluk", "7", indent_cm=0.75)
    add_toc_entry(doc, "BÖLÜM 3: YÖNTEM", "8")
    add_toc_entry(doc, "3.1. Araştırma Modeli", "8", indent_cm=0.75)
    add_toc_entry(doc, "3.2. Evren ve Örneklem", "8", indent_cm=0.75)
    add_toc_entry(doc, "3.3. Yazılım Mimarisi ve Teknoloji Yığını", "9", indent_cm=0.75)
    add_toc_entry(doc, "3.4. Veri Giriş Modülleri", "9", indent_cm=0.75)
    add_toc_entry(doc, "3.5. Genetik Algoritma Motoru", "10", indent_cm=0.75)
    add_toc_entry(doc, "3.6. Puanlama Metrikleri", "10", indent_cm=0.75)
    add_toc_entry(doc, "3.7. Tamamlayıcı Modüller", "11", indent_cm=0.75)
    add_toc_entry(doc, "3.8. Proje İş-Zaman Çizelgesi", "11", indent_cm=0.75)
    add_toc_entry(doc, "3.9. Veri Toplama ve Ölçme Protokolü", "12", indent_cm=0.75)
    add_toc_entry(doc, "3.10. Tekrarlanabilir Benchmark Komutu", "12", indent_cm=0.75)
    add_toc_entry(doc, "Tablo 1. Yazılım teknoloji yığını", "13", indent_cm=0.75)
    add_toc_entry(doc, "Tablo 2. Proje iş-zaman çizelgesi", "13", indent_cm=0.75)
    add_toc_entry(doc, "BÖLÜM 4: BULGULAR VE YORUM", "14")
    add_toc_entry(doc, "4.1. Genetik Algoritma Performansı", "14", indent_cm=0.75)
    add_toc_entry(doc, "4.2. NLP Modülü", "15", indent_cm=0.75)
    add_toc_entry(doc, "4.3. Performans ve Kullanılabilirlik", "15", indent_cm=0.75)
    add_toc_entry(doc, "Tablo 3. Demo seti benchmark sonuçları", "16", indent_cm=0.75)
    add_toc_entry(doc, "Tablo 4. test_students.csv benchmark sonuçları", "16", indent_cm=0.75)
    add_toc_entry(doc, "BÖLÜM 5: SONUÇ VE TARTIŞMA", "17")
    add_toc_entry(doc, "5.1. Sonuç", "17", indent_cm=0.75)
    add_toc_entry(doc, "5.2. Tartışma", "17", indent_cm=0.75)
    add_toc_entry(doc, "5.3. Öneriler", "18", indent_cm=0.75)
    add_toc_entry(doc, "KAYNAKLAR", "19")
    add_toc_entry(doc, "EKLER", "20")
    add_toc_entry(doc, "EK-A: Uygulama Ekran Görüntüleri", "20", indent_cm=0.75)
    add_toc_entry(doc, "EK-B: Örnek Öğrenci Test Veri Seti", "26", indent_cm=0.75)
    add_toc_entry(doc, "EK-C: Tekrarlanabilir Benchmark Betiği", "27", indent_cm=0.75)
    add_toc_entry(doc, "ÖZGEÇMİŞ", "28")


def add_figure(doc, caption, image_path, width_cm=14.5):
    path = Path(image_path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
    else:
        add_body(doc, f"[Ekran görüntüsü bulunamadı: {path.name}]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = cap.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_report(output_path):
    doc = Document()
    style_normal(doc)
    for section in doc.sections:
        set_margins(section)

    # İÇ KAPAK
    add_center_bold(doc, "AKLISIRA: YAPAY ZEKA DESTEKLİ AKILLI SINIF OTURMA DÜZENİ ÖNERİCİ", 14, 72, 36)
    add_center_bold(doc, "Yusuf Kerim KAYMAKCI", 14, 0, 36)
    add_center_bold(doc, "GENEL ZİHİNSEL — BİLİŞİM TEKNOLOJİLERİ YETENEK ALANI", 14, 0, 36)
    add_center_bold(doc, "ESENYURT BİLİM VE SANAT MERKEZİ", 14, 0, 36)
    add_center_bold(doc, "HAZİRAN, 2026", 14, 0, 36)
    add_page_break(doc)

    # ETİK BEYAN
    add_center_bold(doc, "ETİK İLKELERE UYGUNLUK BEYANI", 14, 36, 24)
    add_body(doc, "Proje yazma sürecinde bilimsel ve etik ilkelere uyduğumu, yararlandığım tüm kaynakları kaynak gösterme ilkelerine uygun olarak kaynakçada belirttiğimi ve bu bölümler dışındaki tüm ifadelerin şahsıma ait olduğunu beyan ederim.")
    doc.add_paragraph()
    add_body(doc, "Yazar Adı Soyadı: Yusuf Kerim KAYMAKCI")
    add_body(doc, "Tarih: ....../....../2026")
    add_body(doc, "İmza: ........................................................")
    add_page_break(doc)

    # JÜRİ ONAY
    add_center_bold(doc, "JÜRİ ONAY SAYFASI", 14, 36, 24)
    add_body(doc, 'Yusuf Kerim KAYMAKCI tarafından hazırlanan "AklıSıra: Yapay Zeka Destekli Akıllı Sınıf Oturma Düzeni Önerici" adlı proje çalışması aşağıdaki jüri tarafından oy birliği / oy çokluğu ile kabul edilmiştir.')
    doc.add_paragraph()
    add_body(doc, "Danışman: Bilişim Öğretmeni Oğuz GÜVEN")
    add_body(doc, "İmza: ........................................................")
    doc.add_paragraph()
    add_body(doc, "Jüri Üyesi: ........................................................")
    add_body(doc, "İmza: ........................................................")
    doc.add_paragraph()
    add_body(doc, "Jüri Üyesi: ........................................................")
    add_body(doc, "İmza: ........................................................")
    doc.add_paragraph()
    add_body(doc, "Proje Savunma Tarihi: ....../....../2026  (GG.AA.YYYY)")
    add_body(doc, "Proje Savunma Yeri: Esenyurt Bilim ve Sanat Merkezi / ................................")
    add_page_break(doc)

    # ÖZ
    add_center_bold(doc, "ÖZ", 14, 0, 12)
    for line in [
        "Proje Başlığı: AklıSıra: Yapay Zeka Destekli Akıllı Sınıf Oturma Düzeni Önerici",
        "Proje Türü: Genel Zihinsel — Bilişim Teknolojileri",
        "Proje Tematik Alanı: Eğitim Teknolojileri / Yapay Zeka",
        "Yazar Adı Soyadı: Yusuf Kerim KAYMAKCI",
        "BİLSEM Adı: Esenyurt Bilim ve Sanat Merkezi",
        "Tarih: Haziran, 2026",
    ]:
        add_body(doc, line, single=True)
    doc.add_paragraph()
    add_body(doc, OZ_TEXT, single=True)
    add_body(doc, "Anahtar kelimeler: Yapay zeka, genetik algoritma, oturma düzeni, eğitim teknolojileri, Next.js", single=True)
    add_body(doc, "Danışman: Oğuz GÜVEN", single=True)
    add_page_break(doc)

    # ABSTRACT
    add_center_bold(doc, "ABSTRACT", 14, 0, 12)
    for line in [
        "Project Title: AklıSıra: AI-Powered Smart Classroom Seating Arrangement Recommender",
        "Author: Yusuf Kerim KAYMAKCI",
        "BİLSEM: Esenyurt Science and Art Center",
        "Date: June, 2026",
    ]:
        add_body(doc, line, single=True)
    doc.add_paragraph()
    add_body(doc, ABSTRACT_TEXT, single=True)
    add_body(doc, "Keywords: Artificial intelligence, genetic algorithm, seating arrangement, educational technology, Next.js", single=True)
    add_page_break(doc)

    # İÇİNDEKİLER — BİLSEM: ön bölüm Romen (i–vii), BÖLÜM 1'den itibaren Arap rakamları
    add_table_of_contents(doc)
    add_page_break(doc)

    # SİMGELER
    add_center_bold(doc, "SİMGELER VE KISALTMALAR LİSTESİ", 14, 0, 12)
    add_table(doc, ["Kısaltma", "Açıklama"], [
        ["AI", "Artificial Intelligence (Yapay Zeka)"],
        ["API", "Application Programming Interface"],
        ["BİLSEM", "Bilim ve Sanat Merkezi"],
        ["CSV", "Comma-Separated Values"],
        ["DEHB", "Dikkat Eksikliği ve Hiperaktivite Bozukluğu"],
        ["GA", "Genetik Algoritma"],
        ["JSON", "JavaScript Object Notation"],
        ["KVKK", "Kişisel Verilerin Korunması Kanunu"],
        ["NLP", "Natural Language Processing (Doğal Dil İşleme)"],
        ["SSR", "Server-Side Rendering"],
        ["ZPD", "Zone of Proximal Development (Yakınsak Gelişim Alanı)"],
    ])
    add_page_break(doc)

    # BÖLÜM 1
    add_heading1(doc, "BÖLÜM 1: GİRİŞ")
    add_section_blocks(doc, BOLUM1)
    add_page_break(doc)

    # BÖLÜM 2 — İLGİLİ ARAŞTIRMALAR
    add_heading1(doc, "BÖLÜM 2: İLGİLİ ARAŞTIRMALAR")
    add_section_blocks(doc, BOLUM2_LITERATUR)
    add_page_break(doc)

    # BÖLÜM 3 — YÖNTEM
    add_heading1(doc, "BÖLÜM 3: YÖNTEM")
    add_section_blocks(doc, BOLUM3_YONTEM)
    add_table(doc, ["Katman", "Teknoloji / Araç"], TEKNO_TABLE, "Tablo 1. Yazılım teknoloji yığını")
    add_table(
        doc,
        ["İşin Tanımı", "Kas", "Ara", "Oca", "Şub", "Mar", "Nis", "May", "Haz"],
        IS_ZAMAN_ROWS,
        "Tablo 2. Proje iş-zaman çizelgesi",
    )
    add_page_break(doc)

    # BÖLÜM 4 — BULGULAR
    add_heading1(doc, "BÖLÜM 4: BULGULAR VE YORUM")
    add_section_blocks(doc, BOLUM4_BULGULAR)
    add_table(
        doc,
        ["Metrik", "Rastgele ort. (%)", "GA ort. (%)", "Fark"],
        BULGU_TABLE_DEMO,
        "Tablo 3. Demo seti (15 öğrenci) — 30 rastgele deneme vs. 10 GA koşusu ortalaması",
    )
    add_table(
        doc,
        ["Metrik", "Rastgele ort. (%)", "GA ort. (%)", "Fark"],
        BULGU_TABLE_CSV,
        "Tablo 4. test_students.csv (23 öğrenci) — aynı protokol",
    )
    add_page_break(doc)

    # BÖLÜM 5 — SONUÇ
    add_heading1(doc, "BÖLÜM 5: SONUÇ VE TARTIŞMA")
    add_section_blocks(doc, BOLUM5_SONUC)
    add_page_break(doc)

    # KAYNAKLAR
    add_heading1(doc, "KAYNAKLAR")
    for ref in KAYNAKLAR:
        p = add_body(doc, ref, single=True)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
    add_page_break(doc)

    # EKLER
    add_heading1(doc, "EKLER")
    add_heading2(doc, "EK-A: AklıSıra Uygulaması Ekran Görüntüleri")
    ek_dir = Path(__file__).resolve().parent.parent / "report-ekler"
    figures = [
        (
            "report-ek-01-nlp-giris.png",
            "Şekil 1. Öğrenci rehberi ve serbest metin / Gemini NLP giriş ekranı (aklisira.com/app)",
        ),
        (
            "report-ek-02-sinif-ayarlari.png",
            "Şekil 2. Sınıf ayarları ve düzen tipi seçimi (grid, paired, u-shape, cluster, chevron)",
        ),
        (
            "report-ek-03-ga-optimizasyon.png",
            "Şekil 3. Genetik algoritma optimizasyon sonuçları ve metrik çubuk grafikleri",
        ),
        (
            "report-ek-04-oturma-duzeni.png",
            "Şekil 4. Tamamlanmış oturma düzeni görselleştirmesi (U-düzen, 15 öğrenci)",
        ),
        (
            "report-ek-05-sinav-modu.png",
            "Şekil 5. Kelebek sınav modu (ExamMode) — çok sınıflı salon dağıtımı",
        ),
        (
            "report-ek-06-ai-analiz.png",
            "Şekil 6. Yapay zeka analiz raporu — demo seti en iyi GA koşusu (%98 genel uyum)",
        ),
    ]
    for filename, caption in figures:
        add_figure(doc, caption, ek_dir / filename)
    add_heading2(doc, "EK-B: Örnek Öğrenci Test Veri Seti")
    add_table(doc, ["Name", "Academic Level", "Behavior", "Special Needs"], [
        ["Ali", "high", "leader", "none"],
        ["Ayşe", "average", "disruptive", "adhd"],
        ["Fatma", "above_average", "quiet", "none"],
        ["Mehmet", "struggling", "follower", "vision"],
        ["Elif", "high", "quiet", "none"],
    ], "Tablo 5. test_students.csv dosyasından örnek kayıtlar")
    add_body(doc, "Dosyanın tam içeriği proje klasöründeki public/test_students.csv dosyasında yer almaktadır (23 kayıt).")
    add_heading2(doc, "EK-C: Tekrarlanabilir Benchmark Betiği")
    add_body(doc, "GA ve NLP performans ölçümleri scripts/run-benchmark.ts betiği ile üretilmiştir. "
             "Proje kök dizininden çalıştırmak için: npm run benchmark — çıktı scripts/benchmark-results.json "
             "dosyasına yazılır.")
    add_page_break(doc)

    # ÖZGEÇMİŞ
    add_heading1(doc, "ÖZGEÇMİŞ")
    add_body(doc, "Yusuf Kerim KAYMAKCI")
    add_body(
        doc,
        "18.02.2010 tarihinde Başakşehir'de doğmuştur. İlköğrenimini Şehit Vedat Barceğci İlkokulu'nda tamamlamış; "
        "hâlen Baykar Fen Lisesi 9. sınıf öğrencisidir. Esenyurt Bilim ve Sanat Merkezi'nde Bilişim Teknolojileri "
        "yetenek alanında çalışmalarını sürdürmektedir.",
    )
    add_body(
        doc,
        "UAV sistem entegrasyonu, otonom robot geliştirme ve proje yönetimi alanlarında aktiftir. İHA takımı Zilant "
        "ve Tasarla Geliştir yarışması GÖKBAY takımının kaptanı olarak görev yapmakta; ESP32, ArduPilot, Python ve "
        "CAD/PCB ile donanım–yazılım entegrasyonu gerçekleştirmektedir. TÜBİTAK 2204 kapsamındaki Derim "
        "(kutup/deep-sea benthic lander) projesinde final aşamasında yer almakta; IAC 2026 ADAPT oturumu için kabul "
        "edilmiş bir araştırma özeti üzerinde çalışmaktadır. Codeavour robosoccer ve TUA Astro Hackathon'da ödül "
        "almış; MEB Robot Caretta Caretta kategorisinde yarışmaktadır. Limitless Makers programında mentör olarak "
        "genç ekiplere rehberlik etmekte, BFLTalks'ta konuşmacı olarak yer almaktadır.",
    )
    add_body(
        doc,
        "Bu birikimini AklıSıra projesinde yapay zeka destekli sınıf yönetimi ve optimizasyon algoritmalarına "
        "taşımış; çalışmayı III. Yapay Zeka ile Eğitim Zirvesi'nde sunmuştur.",
    )

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build_report("/home/ykk/Projects/aklisira/Yusuf Kerim Kaymakci-AKLISIRA.docx")
